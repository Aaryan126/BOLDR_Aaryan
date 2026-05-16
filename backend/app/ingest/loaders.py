from __future__ import annotations

import csv
import re
from collections import Counter
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader

from app.models.dataset import (
    DatasetDiagnostics,
    DatasetSamples,
    DatasetSnapshot,
    DocumentChunk,
    DocumentSection,
    ProductModel,
    RateCardItem,
    SourceFile,
    SourceType,
    StrapItem,
    TicketRecord,
)

REPO_ROOT = Path(__file__).resolve().parents[3]
BACKEND_ROOT = Path(__file__).resolve().parents[2]
DATA_ROOT = REPO_ROOT / "Boldr Data"
PROCESSED_DATA_ROOT = BACKEND_ROOT / "data" / "processed"

EXPECTED_BRIEF_FILE_COUNT = 11

SOURCE_SPECS: tuple[dict[str, object], ...] = (
    {
        "file_name": "01_customer_tickets.csv",
        "source_type": "tickets",
        "source_priority": 5,
        "role": "Primary sample inbox and evaluation set",
    },
    {
        "file_name": "03a_rate_card_engraving.csv",
        "source_type": "rate_card",
        "source_priority": 1,
        "role": "Authoritative engraving prices, limits, scripts, and timing",
    },
    {
        "file_name": "03b_rate_card_servicing.csv",
        "source_type": "rate_card",
        "source_priority": 1,
        "role": "Authoritative servicing prices, inclusions, and turnaround",
    },
    {
        "file_name": "04_faq_document.pdf",
        "source_type": "faq",
        "source_priority": 3,
        "role": "Current customer-facing FAQ and reply style source",
    },
    {
        "file_name": "05a_SOP.docx",
        "source_type": "sop",
        "source_priority": 4,
        "role": "Internal workflow, escalation, tone, and routing source",
    },
    {
        "file_name": "05b_product_reference.docx",
        "source_type": "product_reference",
        "source_priority": 2,
        "role": "Authoritative product specs, straps, safety, and availability",
    },
)

FAQ_HEADINGS = [
    "Materials & Safety",
    "Engraving",
    "Strap Compatibility",
    "Watch Servicing",
    "Orders & Shipping",
    "General",
]

SOP_HEADING_PATTERN = re.compile(r"^(?:\d+\.|4\.\d)\s+.+$")
PRODUCT_HEADINGS = [
    "Current Models",
    "Strap Catalogue - All Models (20mm)",
    "Common Questions - Quick Answers",
]
PRODUCT_FIELD_NAMES = {
    "Case",
    "Crystal",
    "Movement",
    "Water resistance",
    "Lume",
    "Weight (no strap)",
    "Dial colours",
    "Dial color",
    "Dial colour",
    "Strap (included)",
    "Strap options",
    "Lug width",
    "Safety",
    "Warranty",
    "Availability",
}


def source_files(data_root: Path = DATA_ROOT) -> list[SourceFile]:
    sources: list[SourceFile] = []
    for spec in SOURCE_SPECS:
        path = data_root / str(spec["file_name"])
        sources.append(
            SourceFile(
                file_name=str(spec["file_name"]),
                relative_path=f"Boldr Data/{spec['file_name']}",
                source_type=spec["source_type"],  # type: ignore[arg-type]
                source_priority=int(spec["source_priority"]),
                role=str(spec["role"]),
                exists=path.exists(),
                size_bytes=path.stat().st_size if path.exists() else None,
            )
        )
    return sources


def load_dataset(data_root: Path = DATA_ROOT) -> DatasetSnapshot:
    sources = source_files(data_root)
    missing = [source.relative_path for source in sources if not source.exists]
    if missing:
        return _empty_snapshot(sources, missing)

    tickets = load_tickets(data_root / "01_customer_tickets.csv")
    engraving_items = load_engraving_rate_card(data_root / "03a_rate_card_engraving.csv")
    servicing_items = load_servicing_rate_card(data_root / "03b_rate_card_servicing.csv")

    faq_sections = load_pdf_sections(
        data_root / "04_faq_document.pdf",
        source_type="faq",
        source_priority=3,
        headings=FAQ_HEADINGS,
    )
    sop_sections = load_sop_sections(data_root / "05a_SOP.docx")
    product_reference_lines = extract_docx_lines(data_root / "05b_product_reference.docx")
    product_sections = load_product_reference_sections(product_reference_lines)
    product_models = parse_product_models(product_reference_lines)
    strap_items = parse_strap_catalogue(product_reference_lines)

    document_sections = [*faq_sections, *sop_sections, *product_sections]
    chunks = chunk_sections(document_sections)

    diagnostics = build_diagnostics(
        sources=sources,
        missing_files=missing,
        tickets=tickets,
        rate_card_items=[*engraving_items, *servicing_items],
        document_sections=document_sections,
        document_chunks=chunks,
        product_models=product_models,
        strap_items=strap_items,
    )

    return DatasetSnapshot(
        sources=sources,
        tickets=tickets,
        rate_card_items=[*engraving_items, *servicing_items],
        document_sections=document_sections,
        document_chunks=chunks,
        product_models=product_models,
        strap_items=strap_items,
        diagnostics=diagnostics,
    )


def load_tickets(path: Path) -> list[TicketRecord]:
    with path.open(newline="", encoding="utf-8-sig") as file:
        rows = list(csv.DictReader(file))

    return [
        TicketRecord(
            ticket_id=row["ticket_id"],
            date_received=row["date_received"],
            customer_name=row["customer_name"],
            customer_email=row["customer_email"],
            order_id=_blank_to_none(row.get("order_id")),
            channel=row["channel"],
            question_type=row["question_type"],
            subject=row["subject"],
            message_body=row["message_body"],
            status=row["status"],
            answered_by_kb=_to_bool(row["answered_by_kb"]),
            requires_escalation=_to_bool(row["requires_escalation"]),
            buyer_persona=row["buyer_persona"],
            agent_notes=_blank_to_none(row.get("agent_notes")),
        )
        for row in rows
    ]


def load_engraving_rate_card(path: Path) -> list[RateCardItem]:
    with path.open(newline="", encoding="utf-8-sig") as file:
        rows = list(csv.DictReader(file))

    return [
        RateCardItem(
            category="engraving",
            source_file=path.name,
            source_priority=1,
            service=_normalize_line(row["service"]),
            price_sgd=_to_float(row.get("price_sgd")),
            notes=_blank_to_none(_normalize_line(row.get("notes", ""))),
        )
        for row in rows
    ]


def load_servicing_rate_card(path: Path) -> list[RateCardItem]:
    with path.open(newline="", encoding="utf-8-sig") as file:
        rows = list(csv.DictReader(file))

    return [
        RateCardItem(
            category="servicing",
            source_file=path.name,
            source_priority=1,
            service=_normalize_line(row["service_tier"]),
            price_sgd=_to_float(row.get("price_sgd")),
            turnaround_days=_blank_to_none(_normalize_line(row.get("turnaround_days", ""))),
            includes=_blank_to_none(_normalize_line(row.get("includes", ""))),
            notes=_blank_to_none(_normalize_line(row.get("notes", ""))),
        )
        for row in rows
    ]


def load_pdf_sections(
    path: Path,
    source_type: SourceType,
    source_priority: int,
    headings: list[str],
) -> list[DocumentSection]:
    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return sections_from_known_headings(
        text=_normalize_text(text),
        source_file=path.name,
        source_type=source_type,
        source_priority=source_priority,
        headings=headings,
    )


def load_sop_sections(path: Path) -> list[DocumentSection]:
    lines = extract_docx_lines(path)
    sections: list[DocumentSection] = []
    current_title = "Preamble"
    current_lines: list[str] = []

    for line in lines:
        if SOP_HEADING_PATTERN.match(line):
            _append_section(
                sections,
                source_file=path.name,
                source_type="sop",
                source_priority=4,
                title=current_title,
                lines=current_lines,
            )
            current_title = line
            current_lines = []
        else:
            current_lines.append(line)

    _append_section(
        sections,
        source_file=path.name,
        source_type="sop",
        source_priority=4,
        title=current_title,
        lines=current_lines,
    )
    return [section.model_copy(update={"order": index}) for index, section in enumerate(sections)]


def load_product_reference_sections(lines: list[str]) -> list[DocumentSection]:
    text = "\n".join(lines)
    normalized = _normalize_text(text).replace("—", "-")
    return sections_from_known_headings(
        text=normalized,
        source_file="05b_product_reference.docx",
        source_type="product_reference",
        source_priority=2,
        headings=PRODUCT_HEADINGS,
    )


def extract_docx_lines(path: Path) -> list[str]:
    document = Document(str(path))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(_normalize_line(text))

    for table in document.tables:
        for row in table.rows:
            cells = [_normalize_line(cell.text) for cell in row.cells]
            non_empty = [cell for cell in cells if cell]
            if non_empty:
                lines.append(" | ".join(non_empty))

    return lines


def sections_from_known_headings(
    text: str,
    source_file: str,
    source_type: SourceType,
    source_priority: int,
    headings: list[str],
) -> list[DocumentSection]:
    matches: list[tuple[int, str]] = []
    for heading in headings:
        index = text.find(heading)
        if index >= 0:
            matches.append((index, heading))
    matches.sort()

    sections: list[DocumentSection] = []
    for order, (start, heading) in enumerate(matches):
        end = matches[order + 1][0] if order + 1 < len(matches) else len(text)
        section_text = text[start + len(heading) : end].strip()
        sections.append(
            DocumentSection(
                source_file=source_file,
                source_type=source_type,
                source_priority=source_priority,
                title=heading,
                text=section_text,
                order=order,
            )
        )
    return sections


def parse_product_models(lines: list[str]) -> list[ProductModel]:
    models: list[ProductModel] = []
    index = 0

    while index < len(lines):
        line = lines[index]
        if "SKU:" not in line or "SGD" not in line:
            index += 1
            continue

        name, price_sgd, sku = _parse_product_model_header(line)
        attributes: dict[str, str] = {}
        index += 1

        while index < len(lines):
            current = lines[index]
            if current.startswith("Strap Catalogue") or ("SKU:" in current and "SGD" in current):
                break
            if current in PRODUCT_FIELD_NAMES and index + 1 < len(lines):
                attributes[current] = lines[index + 1]
                index += 2
            else:
                index += 1

        models.append(ProductModel(name=name, price_sgd=price_sgd, sku=sku, attributes=attributes))

    return models


def parse_strap_catalogue(lines: list[str]) -> list[StrapItem]:
    try:
        start = next(index for index, line in enumerate(lines) if line.startswith("Strap Catalogue"))
    except StopIteration:
        return []

    catalogue_lines = lines[start + 1 :]
    skus = [index for index, line in enumerate(catalogue_lines) if line.startswith("STR-")]
    straps: list[StrapItem] = []

    for sku_index in skus:
        row = [part.strip() for part in catalogue_lines[sku_index].split("|")]
        if len(row) < 6:
            row = catalogue_lines[sku_index : sku_index + 6]
        if len(row) < 6:
            continue
        sku, strap_type, colour, bpa_free, price, compatible_with = row
        straps.append(
            StrapItem(
                sku=sku,
                strap_type=strap_type,
                colour=colour,
                bpa_free=_to_bool(bpa_free),
                price_sgd=_to_float(price),
                compatible_with=compatible_with,
            )
        )

    return straps


def chunk_sections(sections: Iterable[DocumentSection], max_chars: int = 1200) -> list[DocumentChunk]:
    chunks: list[DocumentChunk] = []
    for section in sections:
        section_chunks = _chunk_text(section.text, max_chars=max_chars)
        for chunk_index, chunk in enumerate(section_chunks):
            chunks.append(
                DocumentChunk(
                    chunk_id=f"{section.source_file}:{section.order}:{chunk_index}",
                    source_file=section.source_file,
                    source_type=section.source_type,
                    source_priority=section.source_priority,
                    section_title=section.title,
                    chunk_index=chunk_index,
                    text=chunk,
                )
            )
    return chunks


def build_diagnostics(
    sources: list[SourceFile],
    missing_files: list[str],
    tickets: list[TicketRecord],
    rate_card_items: list[RateCardItem],
    document_sections: list[DocumentSection],
    document_chunks: list[DocumentChunk],
    product_models: list[ProductModel],
    strap_items: list[StrapItem],
) -> DatasetDiagnostics:
    warning = None
    actual_count = sum(1 for source in sources if source.exists)
    if actual_count != EXPECTED_BRIEF_FILE_COUNT:
        warning = (
            f"Challenge brief mentions {EXPECTED_BRIEF_FILE_COUNT} files, "
            f"but {actual_count} actual local data files are available."
        )

    return DatasetDiagnostics(
        expected_brief_file_count=EXPECTED_BRIEF_FILE_COUNT,
        actual_source_file_count=actual_count,
        warning=warning,
        missing_files=missing_files,
        ticket_count=len(tickets),
        answered_by_kb_counts=_bool_counter(ticket.answered_by_kb for ticket in tickets),
        requires_escalation_counts=_bool_counter(ticket.requires_escalation for ticket in tickets),
        question_type_counts=dict(Counter(ticket.question_type for ticket in tickets)),
        buyer_persona_counts=dict(Counter(ticket.buyer_persona for ticket in tickets)),
        channel_counts=dict(Counter(ticket.channel for ticket in tickets)),
        status_counts=dict(Counter(ticket.status for ticket in tickets)),
        engraving_rate_card_count=sum(1 for item in rate_card_items if item.category == "engraving"),
        servicing_rate_card_count=sum(1 for item in rate_card_items if item.category == "servicing"),
        document_section_count=len(document_sections),
        document_chunk_count=len(document_chunks),
        faq_entry_count=sum(
            section.text.count("Q:") for section in document_sections if section.source_type == "faq"
        ),
        product_model_count=len(product_models),
        strap_item_count=len(strap_items),
        source_priorities={source.file_name: source.source_priority for source in sources},
    )


def dataset_samples(snapshot: DatasetSnapshot) -> DatasetSamples:
    engraving_item = next(
        (item for item in snapshot.rate_card_items if item.category == "engraving"), None
    )
    servicing_item = next(
        (item for item in snapshot.rate_card_items if item.category == "servicing"), None
    )

    return DatasetSamples(
        ticket=snapshot.tickets[0] if snapshot.tickets else None,
        engraving_rate_card_item=engraving_item,
        servicing_rate_card_item=servicing_item,
        faq_sections=[
            section.title for section in snapshot.document_sections if section.source_type == "faq"
        ],
        sop_sections=[
            section.title for section in snapshot.document_sections if section.source_type == "sop"
        ],
        product_reference_sections=[
            section.title
            for section in snapshot.document_sections
            if section.source_type == "product_reference"
        ],
        product_models=snapshot.product_models,
        strap_items=snapshot.strap_items[:3],
    )


def write_snapshot(snapshot: DatasetSnapshot, output_path: Path | None = None) -> Path:
    output_path = output_path or PROCESSED_DATA_ROOT / "dataset_snapshot.json"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(snapshot.model_dump_json(indent=2), encoding="utf-8")
    return output_path


def _empty_snapshot(sources: list[SourceFile], missing: list[str]) -> DatasetSnapshot:
    diagnostics = DatasetDiagnostics(
        expected_brief_file_count=EXPECTED_BRIEF_FILE_COUNT,
        actual_source_file_count=sum(1 for source in sources if source.exists),
        warning="One or more required local dataset files are missing.",
        missing_files=missing,
        ticket_count=0,
        answered_by_kb_counts={},
        requires_escalation_counts={},
        question_type_counts={},
        buyer_persona_counts={},
        channel_counts={},
        status_counts={},
        engraving_rate_card_count=0,
        servicing_rate_card_count=0,
        document_section_count=0,
        document_chunk_count=0,
        faq_entry_count=0,
        product_model_count=0,
        strap_item_count=0,
        source_priorities={source.file_name: source.source_priority for source in sources},
    )
    return DatasetSnapshot(
        sources=sources,
        tickets=[],
        rate_card_items=[],
        document_sections=[],
        document_chunks=[],
        product_models=[],
        strap_items=[],
        diagnostics=diagnostics,
    )


def _append_section(
    sections: list[DocumentSection],
    source_file: str,
    source_type: SourceType,
    source_priority: int,
    title: str,
    lines: list[str],
) -> None:
    text = "\n".join(line for line in lines if line.strip()).strip()
    if not text:
        return
    sections.append(
        DocumentSection(
            source_file=source_file,
            source_type=source_type,
            source_priority=source_priority,
            title=title,
            text=text,
            order=len(sections),
        )
    )


def _parse_product_model_header(line: str) -> tuple[str, float | None, str | None]:
    parts = [part.strip() for part in line.split("|")]
    name = parts[0]
    price = None
    sku = None
    for part in parts[1:]:
        if part.startswith("SGD"):
            price = _to_float(part)
        elif part.startswith("SKU:"):
            sku = part.replace("SKU:", "", 1).strip()
    return name, price, sku


def _chunk_text(text: str, max_chars: int) -> list[str]:
    paragraphs = [paragraph.strip() for paragraph in text.split("\n") if paragraph.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        if not current:
            current = paragraph
            continue
        if len(current) + len(paragraph) + 1 <= max_chars:
            current = f"{current}\n{paragraph}"
        else:
            chunks.append(current)
            current = paragraph

    if current:
        chunks.append(current)
    return chunks


def _bool_counter(values: Iterable[bool]) -> dict[str, int]:
    counter = Counter("yes" if value else "no" for value in values)
    return dict(counter)


def _blank_to_none(value: str | None) -> str | None:
    if value is None:
        return None
    value = value.strip()
    return value or None


def _to_bool(value: str | bool | None) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    return value.strip().lower() in {"yes", "true", "1", "y"}


def _to_float(value: str | None) -> float | None:
    if value is None:
        return None
    match = re.search(r"\d+(?:\.\d+)?", value.replace(",", ""))
    return float(match.group(0)) if match else None


def _normalize_line(text: str) -> str:
    return _normalize_text(text).replace("—", "-").replace("–", "-").strip()


def _normalize_text(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", text.replace("\xa0", " ")).strip()
